import os
import fitz  # PyMuPDF for PDF text extraction
import docx
import torch
from fastapi import FastAPI, Depends, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from transformers import AutoTokenizer, AutoModelForSequenceClassification
from sqlalchemy.orm import Session
from .database import get_db, engine, Base, SessionLocal
from .models import Document

app = FastAPI()

# Initialize the database tables
Base.metadata.create_all(bind=engine)

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploaded_files"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Define allowed file extensions
ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}

# Load fine-tuned model and tokenizer
MODEL_PATH = "./fine_tuned_bart"  # Ensure your fine-tuned model is saved here
tokenizer = AutoTokenizer.from_pretrained(MODEL_PATH)
model = AutoModelForSequenceClassification.from_pretrained(MODEL_PATH)
model.eval()  # Set model to evaluation mode

# Define the same categories used during fine-tuning
CATEGORIES = [
    "Technical Documentation",
    "Business Proposal",
    "Legal Document",
    "Academic Paper",
    "General Article",
    "Other",
]

# Create mapping for label indices
category_mapping = {i: label for i, label in enumerate(CATEGORIES)}

def extract_text(file_path: str) -> str:
    """Extract text from a given file based on its extension."""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            doc = fitz.open(file_path)
            return "\n".join([page.get_text() for page in doc])
        elif ext == ".docx":
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        else:
            return None  # Unsupported file format
    except Exception as e:
        return None  # Error in processing

@app.post("/upload/")
async def upload_and_classify_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Handles file uploads, validates file type, extracts text, classifies it using the fine-tuned model, and saves results."""
    
    ext = os.path.splitext(file.filename)[1].lower()
    print(f"📂 Uploaded File Extension: {ext}")  # Debugging output in logs

    # Validate file extension before saving
    if ext not in ALLOWED_EXTENSIONS:
        return JSONResponse(
            status_code=400, 
            content={"error": f"❌ Unsupported file type '{ext}'. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"}
        )

    # Save file only if it is valid
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as f:
        f.write(await file.read())

    # Extract text from the file
    text = extract_text(file_path)
    if not text:
        return JSONResponse(status_code=400, content={"error": "❌ Failed to extract text from the file. It may be empty or corrupted."})

    # Tokenize text using fine-tuned model tokenizer
    inputs = tokenizer(text, padding="max_length", truncation=True, return_tensors="pt")

    # Perform inference with the fine-tuned model
    with torch.no_grad():
        outputs = model(**inputs)
        logits = outputs.logits
        predicted_label_idx = torch.argmax(logits, dim=-1).item()

    # Get category label and confidence
    top_category = category_mapping[predicted_label_idx]
    confidence = torch.softmax(logits, dim=-1)[0][predicted_label_idx].item()

    # Store file and classification details in the database
    document = Document(
        filename=file.filename,
        file_path=file_path,
        category=top_category,
        confidence=confidence,
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    return {
        "message": "✅ File uploaded and classified successfully",
        "file_id": document.id,
        "file_path": document.file_path,
        "category": document.category,
        "confidence": document.confidence,
    }
